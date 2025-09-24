import os
from PyPDF2 import PdfReader
import docx
from io import BytesIO

class DocumentProcessor:
    def __init__(self):
        self.supported_formats = ['txt', 'pdf', 'docx']
    
    def extract_text_from_pdf(self, file_path_or_bytes):
        """Extract text from PDF file"""
        try:
            if isinstance(file_path_or_bytes, str):
                # File path
                reader = PdfReader(file_path_or_bytes)
            else:
                # Bytes object (for uploaded files)
                reader = PdfReader(BytesIO(file_path_or_bytes))
            
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            
            return text
        except Exception as e:
            print(f"Error reading PDF: {e}")
            return ""
    
    def extract_text_from_docx(self, file_path_or_bytes):
        """Extract text from DOCX file"""
        try:
            if isinstance(file_path_or_bytes, str):
                # File path
                doc = docx.Document(file_path_or_bytes)
            else:
                # Bytes object (for uploaded files)
                doc = docx.Document(BytesIO(file_path_or_bytes))
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            return text
        except Exception as e:
            print(f"Error reading DOCX: {e}")
            return ""
    
    def extract_text_from_txt(self, file_path_or_bytes):
        """Extract text from TXT file"""
        try:
            if isinstance(file_path_or_bytes, str):
                # File path
                with open(file_path_or_bytes, 'r', encoding='utf-8') as f:
                    text = f.read()
            else:
                # Bytes object (for uploaded files)
                text = file_path_or_bytes.decode('utf-8')
            
            return text
        except Exception as e:
            print(f"Error reading TXT: {e}")
            return ""
    
    def get_file_type(self, filename):
        """Get file extension"""
        return filename.split('.')[-1].lower()
    
    def process_document(self, file_path_or_bytes, filename=None):
        """Main document processing method"""
        if isinstance(file_path_or_bytes, str):
            # File path provided
            if not os.path.exists(file_path_or_bytes):
                return "Error: File not found"
            file_type = self.get_file_type(file_path_or_bytes)
        else:
            # Bytes object with filename
            if not filename:
                return "Error: Filename required for bytes input"
            file_type = self.get_file_type(filename)
        
        if file_type not in self.supported_formats:
            return f"Error: Unsupported file format. Supported: {', '.join(self.supported_formats)}"
        
        # Process based on file type
        if file_type == 'pdf':
            text = self.extract_text_from_pdf(file_path_or_bytes)
        elif file_type == 'docx':
            text = self.extract_text_from_docx(file_path_or_bytes)
        elif file_type == 'txt':
            text = self.extract_text_from_txt(file_path_or_bytes)
        
        return text if text.strip() else "Error: No text found in document"

# Test the processor
if __name__ == "__main__":
    processor = DocumentProcessor()
    
    # Test with a simple text file
    test_text = "Hello world! This is a test document."
    with open("test_document.txt", "w") as f:
        f.write(test_text)
    
    result = processor.process_document("test_document.txt")
    print("Extracted text:", result)
    
    # Clean up
    os.remove("test_document.txt")